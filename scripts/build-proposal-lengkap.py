# -*- coding: utf-8 -*-
"""
Generate PROPOSAL.docx LENGKAP — satu file .docx sesuai aplikasi SPKT Digital.

Cara pakai:
    python scripts/build-proposal-lengkap.py

Output:
    PROPOSAL.docx  (di root project)

Strategi:
    1. Salin PROPOSAL_backup.docx (pertahankan gambar, tabel literature, format asli)
    2. Patch semua paragraf & tabel agar selaras SPKT Digital (Next.js, SQLite, modul lengkap)
    3. Tambah lampiran teknis aplikasi
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
BACKUP = ROOT / "PROPOSAL_backup.docx"
OUT = ROOT / "PROPOSAL.docx"


# ---------------------------------------------------------------------------
# Penggantian teks paragraf (old -> new)
# ---------------------------------------------------------------------------
PARA_REPLACEMENTS: dict[str, str] = {
    # Typo & kesalahan
    "latara belakang": "latar belakang",
    "masyrakat": "masyarakat",
    "Operasioanl (Portability)": "Operasional (Portability)",
    "BAB IPENDAHULUAN": "BAB I PENDAHULUAN",
    "BAB IIIMETODOLOGI": "BAB III METODOLOGI",
    "BAB IVHASIL": "BAB IV HASIL",
    "HOPOTESIS": "HIPOTESIS",
    "sistem pakar": "sistem informasi SPKT Digital",
    "Sistem Pakar": "Sistem Informasi SPKT Digital",
    # BAB I — tambahan konteks aplikasi
    "Berdasarkan situasi tersebut, dibutuhkan sebuah pengembangan sistem informasi SPKT yang berbasis web, yang memungkin warga untuk melaporkan secara daring tanpa harus ke kantor secara langsung. Sistem ini diharapkan mampu menawarkan fitur pengisian data laporan secara mandiri oleh masyarakat, penyimpanan data secara terpusat, serta memberikan kemudahan bagi petugas dalam mengelola dan mengawasi data laporan yang diterima.":
    "Berdasarkan situasi tersebut, penelitian ini membangun SPKT Digital — sistem informasi berbasis web (Next.js 15, React 19, SQLite) yang memungkinkan masyarakat melapor, mengajukan surat, pengaduan, dan melacak status layanan secara daring. Petugas memiliki antrian digital (officer inbox); administrator mengelola user, statistik, dan audit log.",
  "Selain itu dalam sistem yang dibuat juga dimasukkan metode Customer Satisfaction Index (CSI) sebagai fitur untuk memproses informasi dari kuesioner yang dimasukkan kedalam sistem. Metode ini digunakan sebagai salah satu bagian dari proses otomatisasi pengolahan data dalam sistem, sehingga mendukung peran sistem sebagai alat bantu dalam manajemen data.":
    "Modul Customer Satisfaction Index (CSI) terintegrasi dalam aplikasi: setelah layanan selesai, pengguna mengisi survei 5 dimensi (skala 1–4) dan sistem menghitung skor CSI otomatis dari database satisfaction_surveys.",
    "Belum tersedianya sistem yang memungkinkan masyarakat untuk melakukan pengisian data laporan secara mandiri melalui media online.":
    "Belum tersedia portal web terintegrasi untuk laporan, surat, pengaduan, pelacakan status, dan survei kepuasan (CSI) secara online.",
    "Penelitian ini difokuskan pada perancangan dan pengembangan sistem informasi Sentra Pelayanan Kepolisian Terpadu (SPKT) berbasis web pada Polsek Palmerah.":
    "Penelitian ini merancang dan membangun SPKT Digital berbasis web pada Polsek Palmerah, meliputi autentikasi 3 role, laporan polisi, layanan surat, pengaduan, pelacakan publik, notifikasi, dan CSI terintegrasi.",
    "Penilaian tingkat kepuasan masyarakat pada penelitian ini dibatasi hanya menggunakan metode Customer Satisfaction Index (CSI) setelah pengguna menggunakan sistem informasi SPKT berbasis web.":
    "Penilaian kepuasan menggunakan metode CSI terintegrasi aplikasi (5 dimensi, skala Likert 1–4) setelah layanan laporan/surat/pengaduan selesai.",
    "Bab ini menyajikan hasil implementasi sistem pakar yang telah dikembangkan, hasil pengujian, analisa perbandingan tingkat akurasi dan efektivitas, serta pembahasan berdasarkan data penelitian yang diperoleh.":
    "Bab ini menyajikan hasil implementasi sistem informasi SPKT Digital, pengujian black-box per role, evaluasi modul CSI, serta pembahasan berdasarkan data penelitian.",
    # 2.8 Kebutuhan fungsional
    "Login Pengguna": "Autentikasi & Keamanan",
    "Sistem menyediakan fitur login bagi petugas SPKT menggunakan username dan password untuk menjaga keamanan akses pengguna.":
    "Registrasi/login email-password (bcrypt) untuk masyarakat, petugas, dan admin; sesi cookie HttpOnly (spkt_session); 2FA TOTP; lupa password; logout; RBAC 3 role.",
    "Input Data Pelapor": "Pelaporan Online (Masyarakat)",
    "Sistem memungkinkan petugas memasukan data pelapor ke dalam sistem.":
    "Masyarakat mengisi laporan sendiri (jenis kejadian, lokasi, deskripsi), upload bukti ke data/uploads/, simpan draft, kirim dengan nomor LAP-YYYY-NNNN dan timeline status.",
    "Pengelolaan Data Pelapor": "Antrian & Pemrosesan Petugas",
    "Sistem dapat digunakan untuk mengelola data laporan masyarakat seperti menambah,":
    "Petugas melihat antrian officer inbox (belum ditugaskan + tugas sendiri), verifikasi, assign (assigned_officer_id), proses, selesaikan/tolak; timeline + notifikasi ke pelapor.",
    "Menyimpan, dan memperbaharui status laporan.": "Admin dapat override status dengan pencatatan audit log.",
    "Pembuatan Surat Tanda Penerimaan Laporan Kepolisian (STPLK)": "Layanan Surat & Pengaduan",
    "Sistem mampu memproses pembuatan Surat Tanda Penerimaan Laporan Kepolisian (STPLK).":
    "Pengajuan surat (SUR-YYYY-NNNN) dan pengaduan (PGD-YYYY-NNNN) online; petugas update status; unduh PDF surat; timeline per layanan.",
    "Cetak Surat Tanda Penerimaan Laporan Kepolisian (STPLK)": "Pelacakan Layanan Publik",
    "Sistem menyediakan fitur untuk mencetak STPLK yang akan diberikan kepada pelapor.":
    "GET /api/track — lacak laporan/surat/pengaduan dengan nomor + NIK tanpa login (rate limit).",
    "Penilaian Kepuasan Masyarakat": "Modul CSI Terintegrasi",
    "Sistem menyediakan fasilitas pengisian kuesioner kepuasan masyarakat berdasarkan metode Customer Satisfaction Index (CSI).":
    "Survei otomatis setelah layanan selesai: 5 dimensi (kemudahan, kecepatan, keramahan petugas, kejelasan informasi, kualitas hasil), skala 1–4.",
    "Perhitungan Nilai CSI": "Dashboard Admin & Notifikasi",
    "Sistem dapat menghitung nilai kepuasan masyarakat secara otomatis menggunakan metode CSI.":
    "Perhitungan CSI otomatis di backend; dashboard admin (/api/stats/admin, /api/survey/csi/summary); notifikasi in-app; kelola user/petugas/artikel.",
    "Sistem mampu menyimpan data pelapor, laporan, dan hasil penilaian kepuasan masyarakat ke dalam database.":
    "Database SQLite data/spkt.db (21 tabel) + file upload data/uploads/. Stack: Next.js 15, React 19, TypeScript, REST API ~49 endpoint.",
    "Sistem berbasis website sehingga dapat diakses melalui perangkat komputer maupun smartphone yang memiliki koneksi internet":
    "Antarmuka responsif (Tailwind CSS + Radix UI), diakses via browser desktop/mobile; Node.js ≥ 22.5; healthcheck /api/health.",
    # 2.9 Desain
    "Perancangan sistem difokuskan pada pembuatan sistem informasi SPKT berbasis website yang mendukung proses pelayanan dan pelaporan masyarakat secara online. Sistem ini terdiri dari Pelapor, petugas SPKT, website pelayanan, dan database penyimpanan data laporan.":
    "SPKT Digital: 3 role (masyarakat/user, petugas, admin), modul laporan/surat/pengaduan/lacak/notifikasi/CSI. Arsitektur 3 lapis: Frontend (React + spktApi.ts) → Backend (Next.js API Routes + services/) → Data (SQLite + uploads).",
    "Pelapor menggunakan sistem SPKT untuk membuat laporan dengan mengisi data yang diperlukan.":
    "Entitas users (pusat) terhubung sessions, notifications, officers, reports, letter_requests, complaints, satisfaction_surveys, audit_logs — total 21 tabel SQLite.",
    "Laporan berfungsi untuk menyimpan data pelaporan masyarakat yang masuk ke sistem.":
    "reports + report_timeline + report_evidence; assigned_officer_id → officers; nomor LAP via reference_counters.",
    "Petugas SPKT bertugas memverifikasi dan memproses laporan yang diterima.":
    "Petugas: officer inbox, verifikasi (submitted→verified), assign, proses (processing), selesai (completed) atau tolak (rejected).",
    "Penilaian CSI digunakan untuk menyimpan data penilaian kepuasan masyarakat berdasarkan indikator yang tersedia.":
    "satisfaction_surveys + survey_responses + survey_dimensions (ease, speed, officer, clarity, quality).",
    "Indikator CSI berfungsi sebagai acuan dalam proses pengukuran kepuasan masyarakat.":
    "Skor per dimensi 1–4; csi_score = rata-rata tertimbang / 4 × 100%.",
    # Hipotesis
    "Diduga bahwa perancangan sistem informasi Sentra Pelayanan Kepolisian Terpadu (SPKT) berbasis web dapat mengubah proses pelaporan masyrakat yang sebelumnya dilakukan secara langsung menjadi dapat dilakukan secara online.":
    "Diduga SPKT Digital dapat mengubah pelaporan manual menjadi online (laporan, surat, pengaduan).",
    "Diduga bahwa Sistem Informasi SPKT berbasis web yang dikembangkan dapat memberikan kemudahan bagi masyrakat dalam melakukan pelaporan tanpa harus datang langsung ke Polsek Palmerah.":
    "Diduga SPKT Digital mempermudah masyarakat melapor, lacak status, dan isi CSI tanpa datang ke kantor.",
    "Diduga bahwa penerapan sistem informasi SPKT berbasis web dapat meningkatkan efektivitas dan efisiensi dalam pengelolaan data laporan masyarakat oleh petugas.":
    "Diduga SPKT Digital meningkatkan efektivitas antrian petugas dan pengelolaan data terpusat.",
    # CSI kuesioner lama -> dimensi aplikasi
    "Kecepatan petugas dalam menerima laporan/pengaduan.": "1. Kemudahan Prosedur (ease)",
    "Kejelasan prosedur dalam pelayanan laporan/pengaduan.": "2. Kecepatan Pelayanan (speed)",
    "Kemudahan masyarakat dalam mengakses pelayanan SPKT.": "3. Keramahan Petugas (officer)",
    "Ketersediaan sarana dan prasarana pelayanan (Ruang tunggu, Kursi, AC, dan Toilet).": "4. Kejelasan Informasi (clarity)",
    "Ketepatan waktu penyelesaian layanan sesuai janji.": "5. Kualitas Hasil Layanan (quality)",
}

# Regex replacements (partial match)
PARA_REGEX: list[tuple[str, str]] = [
    (r"memungkin warga", "memungkinkan warga"),
]

# ---------------------------------------------------------------------------
# Update tabel
# ---------------------------------------------------------------------------
TABLE_7_USULAN = [
    ("Komponen", "Deskripsi"),
    ("Nama use case", "Layanan SPKT Digital (laporan, surat, pengaduan, lacak, CSI)"),
    ("Actor Utama", "Masyarakat, Petugas SPKT, Admin, Pengunjung anonim"),
    ("Deskripsi", "Portal web SPKT Digital: register/login, buat laporan & draft, ajukan surat/pengaduan, antrian petugas, update status, notifikasi, survei CSI, kelola admin."),
    ("Kondisi Awal", "Pengguna membuka http://localhost:3000 atau URL deploy."),
    ("Kondisi Akhir", "Layanan diproses, status & timeline terupdate, CSI tersimpan, notifikasi terkirim."),
]

TABLE_8_CRC = [
    ("No", "Kelas", "Responsibility (Tanggung Jawab)", "Collaboration (Relasi)"),
    ("1.", "User", "Login, logout, preferensi notifikasi/privasi", "Report, Letter, Complaint"),
    ("2.", "Report", "Simpan laporan, timeline, evidence, status", "User, Officer"),
    ("3.", "LetterRequest", "Ajukan surat, lampiran, PDF", "User, Officer"),
    ("4.", "Complaint", "Ajukan pengaduan, tanggapan", "User, Officer"),
    ("5.", "Officer", "Verifikasi, assign, proses layanan", "Report, Letter, Complaint"),
    ("6.", "Admin", "Override status, kelola user/petugas, audit", "User, Officer, AuditLog"),
    ("7.", "SatisfactionSurvey", "Simpan skor CSI 5 dimensi", "SurveyDimension, User"),
]

TABLE_9_LIKERT = [
    ("No", "Kriteria Jawaban", "Skor"),
    ("1.", "Sangat Tidak Puas", "1"),
    ("2.", "Tidak Puas", "2"),
    ("3.", "Puas", "3"),
    ("4.", "Sangat Puas", "4"),
]

LAMPIRAN_ROWS = [
    ("Framework", "Next.js 15 (App Router)"),
    ("Frontend", "React 19, TypeScript, Tailwind CSS, Radix UI"),
    ("Client API", "src/lib/spktApi.ts"),
    ("Backend", "Next.js API Routes (~49 endpoint REST)"),
    ("Services", "spkt.ts, users.ts, notifications.ts, audit.ts, csi.ts"),
    ("Database", "SQLite data/spkt.db — 21 tabel"),
    ("Upload", "data/uploads/"),
    ("Auth", "Cookie spkt_session, bcrypt, 2FA TOTP"),
    ("Role", "user (masyarakat), petugas, admin"),
    ("Nomor laporan", "LAP-YYYY-NNNN"),
    ("Nomor surat", "SUR-YYYY-NNNN"),
    ("Nomor pengaduan", "PGD-YYYY-NNNN"),
    ("Akun demo user", "user@spkt.id / spkt123"),
    ("Akun demo petugas", "petugas@spkt.id / spkt123"),
    ("Akun demo admin", "admin@spkt.id / spkt123"),
    ("Jalankan lokal", "npm install && npm run dev → http://localhost:3000"),
    ("Dokumentasi", "docs/presentasi-ppt-teknis.md, README.md"),
]


def replace_in_paragraph(para, mapping: dict[str, str]) -> bool:
    text = para.text
    if not text.strip():
        return False
    new = text
    for old, repl in mapping.items():
        if old in new:
            new = new.replace(old, repl)
    for pattern, repl in PARA_REGEX:
        new = re.sub(pattern, repl, new)
    if new != text:
        if para.runs:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = new
        return True
    return False


def patch_all_paragraphs(doc: Document) -> int:
    count = 0
    for para in doc.paragraphs:
        if replace_in_paragraph(para, PARA_REPLACEMENTS):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph(para, PARA_REPLACEMENTS):
                        count += 1
    return count


def set_table_rows(table, rows_data: list[tuple]) -> None:
    """Resize table if needed and fill cells."""
    needed = len(rows_data)
    cols = len(rows_data[0])
    while len(table.rows) < needed:
        table.add_row()
    while len(table.rows) > needed:
        tbl = table._tbl
        tbl.remove(table.rows[-1]._tr)
    for ri, row_data in enumerate(rows_data):
        for ci in range(cols):
            table.rows[ri].cells[ci].text = str(row_data[ci])


def patch_tables(doc: Document) -> None:
    if len(doc.tables) > 7:
        set_table_rows(doc.tables[7], TABLE_7_USULAN)
    if len(doc.tables) > 8:
        set_table_rows(doc.tables[8], TABLE_8_CRC)
    if len(doc.tables) > 9:
        set_table_rows(doc.tables[9], TABLE_9_LIKERT)


def trim_extra_questionnaire_paragraphs(doc: Document) -> None:
    """Hapus pertanyaan kuesioner lama (item 6–15) setelah 5 dimensi CSI."""
    skip_starts = (
        "Kesopanan dan keramahan",
        "Kemampuan petugas",
        "Tindak lanjut",
        "Keamanan dan kenyamanan",
        "Kebersihan ruang",
        "Ketersediaan tempat parkir",
        "Ketersediaan media penyampaian",
        "Ketepatan jam operasional",
        "Keadilan dan ketidakberpihakan",
    )
    to_remove = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if any(t.startswith(s) for s in skip_starts):
            to_remove.append(para)
    for para in to_remove:
        el = para._element
        el.getparent().remove(el)


def append_lampiran_teknis(doc: Document) -> None:
    """Tambah lampiran spesifikasi aplikasi di akhir dokumen."""
    doc.add_page_break()
    doc.add_heading("LAMPIRAN A — SPESIFIKASI TEKNIS SPKT DIGITAL", level=1)
    doc.add_paragraph(
        "Lampiran ini merangkum implementasi aktual aplikasi yang dibangun, "
        "selaras dengan kebutuhan fungsional pada Bab II dan kode sumber di repositori proyek."
    )
    t = doc.add_table(rows=1 + len(LAMPIRAN_ROWS), cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Komponen"
    t.rows[0].cells[1].text = "Spesifikasi"
    for run in t.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True
    for run in t.rows[0].cells[1].paragraphs[0].runs:
        run.bold = True
    for i, (k, v) in enumerate(LAMPIRAN_ROWS, start=1):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v

    doc.add_heading("LAMPIRAN B — Alur Status Layanan", level=2)
    for line in [
        "Laporan: draft → submitted → verified → assigned → processing → completed | rejected",
        "Surat: draft → submitted → verified → ready → completed | rejected",
        "Pengaduan: submitted → reviewing → processing → resolved → closed",
    ]:
        doc.add_paragraph("• " + line)

    doc.add_heading("LAMPIRAN C — Modul API Utama", level=2)
    apis = [
        "POST /api/auth/login, GET /api/auth/session, POST /api/auth/register",
        "GET/POST /api/reports, PATCH /api/reports/:id",
        "GET/POST /api/letters, GET /api/letters/:id/pdf",
        "GET/POST /api/complaints, PATCH /api/complaints/:id",
        "GET /api/track (publik), POST /api/upload",
        "POST /api/survey/submit, GET /api/survey/csi/summary",
        "GET /api/stats/admin, GET /api/audit-logs",
    ]
    for a in apis:
        doc.add_paragraph("• " + a)


def main() -> None:
    if not BACKUP.exists():
        raise FileNotFoundError(f"File backup tidak ditemukan: {BACKUP}")

    shutil.copy2(BACKUP, OUT)
    doc = Document(str(OUT))

    n = patch_all_paragraphs(doc)
    patch_tables(doc)
    trim_extra_questionnaire_paragraphs(doc)
    append_lampiran_teknis(doc)

    doc.save(str(OUT))
    print(f"Selesai: {OUT}")
    print(f"  - Paragraf di-patch: {n}")
    print(f"  - Gambar/diagram: dipertahankan dari backup")
    print(f"  - Tabel literature review: dipertahankan")
    print(f"  - Lampiran teknis A/B/C: ditambahkan")
    print("\nBuka PROPOSAL.docx -> References -> Update Table of Contents")


if __name__ == "__main__":
    main()
