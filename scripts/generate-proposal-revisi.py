# -*- coding: utf-8 -*-
"""Generate PROPOSAL.docx aligned with SPKT Digital application."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"c:\Users\najri\Downloads\Digital Police Service Website\PROPOSAL.docx"


def add_title(doc, text, size=14, bold=True, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # === COVER ===
    add_title(doc, "RANCANG BANGUN SISTEM INFORMASI SENTRA PELAYANAN KEPOLISIAN TERPADU (SPKT) DENGAN METODE CUSTOMER SATISFACTION INDEX (CSI) PADA POLISI SEKTOR PALMERAH", 13)
    doc.add_paragraph()
    add_title(doc, "PROPOSAL SKRIPSI", 12)
    doc.add_paragraph()
    add_body(doc, "Disusun Oleh:")
    add_body(doc, "Nama Mahasiswa\t: Sitriani Masu")
    add_body(doc, "NIM\t\t: 22411016")
    add_body(doc, "Fakultas\t\t: Teknologi")
    add_body(doc, "Program Studi\t: Sistem Informasi")
    add_body(doc, "Jenjang\t\t: S-1 (Strata Satu)")
    doc.add_paragraph()
    add_title(doc, "FAKULTAS TEKNOLOGI")
    add_title(doc, "INSTITUT SOSIAL DAN TEKNOLOGI WIDURI")
    add_title(doc, "JAKARTA")
    add_title(doc, "2026")
    doc.add_page_break()

    # === KATA PENGANTAR ===
    add_heading(doc, "KATA PENGANTAR", 1)
    add_body(doc, (
        "Puji dan syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa sehingga proposal skripsi "
        "berjudul \"Rancang Bangun Sistem Informasi Sentra Pelayanan Kepolisian Terpadu (SPKT) "
        "Dengan Metode Customer Satisfaction Index (CSI) Pada Polisi Sektor Palmerah\" dapat diselesaikan."
    ))
    add_body(doc, (
        "Penelitian ini dilaksanakan di SPKT Polsek Palmerah Jakarta Barat. Penulis mengucapkan terima kasih "
        "kepada dosen pembimbing Elizabeth N. Tamatjita, S.Kom., M.M., M.Cs., M.H., petugas SPKT, "
        "serta keluarga dan rekan yang telah memberikan dukungan."
    ))
    add_body(doc, "Jakarta, 06 Juni 2026\nSitriani Masu")
    doc.add_page_break()

    # === DAFTAR ISI placeholder ===
    add_heading(doc, "DAFTAR ISI", 1)
    add_body(doc, "(Perbarui daftar isi otomatis di Microsoft Word: References → Table of Contents)")
    doc.add_page_break()

    # === BAB I ===
    add_heading(doc, "BAB I PENDAHULUAN", 1)
    add_heading(doc, "1.1 Latar Belakang", 2)
    add_body(doc, (
        "Pelayanan publik merupakan kegiatan lembaga pemerintahan untuk memenuhi kebutuhan masyarakat. "
        "Di Indonesia, Kepolisian Negara Republik Indonesia (Polri) bertanggung jawab menjaga keamanan dan ketertiban. "
        "Polisi Sektor Palmerah Jakarta Barat melayani masyarakat melalui Sentra Pelayanan Kepolisian Terpadu (SPKT) "
        "sebagai garda terdepan penerima laporan dan pengaduan."
    ))
    add_body(doc, (
        "SPKT Polsek Palmerah saat ini telah memanfaatkan komputerisasi internal, namun masyarakat masih harus "
        "datang langsung untuk pelaporan. Belum tersedia portal web terintegrasi yang memungkinkan masyarakat "
        "mengajukan laporan, surat, dan pengaduan secara daring serta memantau status layanan secara mandiri."
    ))
    add_body(doc, (
        "Berdasarkan kondisi tersebut, penelitian ini merancang dan membangun Sistem Informasi SPKT berbasis web "
        "(SPKT Digital) menggunakan Next.js 15, React, dan SQLite. Sistem mencakup tiga peran pengguna "
        "(masyarakat, petugas, administrator), layanan laporan polisi, layanan surat, pengaduan, pelacakan layanan, "
        "notifikasi in-app, dan pengukuran kepuasan masyarakat dengan metode Customer Satisfaction Index (CSI) "
        "terintegrasi dalam aplikasi."
    ))

    add_heading(doc, "1.2 Identifikasi Masalah", 2)
    add_bullets(doc, [
        "Pelayanan SPKT Polsek Palmerah belum menyediakan sistem berbasis web yang dapat diakses masyarakat secara online.",
        "Proses pelaporan masih bergantung pada kedatangan fisik sehingga waktu tunggu dan antrian menjadi kendala.",
        "Belum ada sistem terpusat untuk pelacakan status laporan, surat, dan pengaduan oleh masyarakat.",
        "Pengukuran kepuasan masyarakat belum terkomputerisasi dalam satu sistem informasi SPKT.",
    ])

    add_heading(doc, "1.3 Rumusan Masalah", 2)
    add_numbered(doc, [
        "Apakah perancangan sistem informasi SPKT berbasis web dapat mengubah proses pelaporan dari manual menjadi online?",
        "Apakah sistem yang dibangun mempermudah masyarakat melapor tanpa harus datang ke Polsek Palmerah?",
        "Bagaimana proses perancangan dan pembangunan sistem informasi SPKT berbasis web pada Polsek Palmerah?",
        "Bagaimana penerapan metode CSI dalam sistem untuk mengukur kepuasan masyarakat terhadap layanan digital?",
    ])

    add_heading(doc, "1.4 Tujuan dan Manfaat Penelitian", 2)
    add_heading(doc, "1.4.1 Tujuan Penelitian", 3)
    add_numbered(doc, [
        "Merancang dan membangun sistem informasi SPKT berbasis web (SPKT Digital) sesuai kebutuhan SPKT Polsek Palmerah.",
        "Menyediakan layanan pelaporan online, layanan surat, pengaduan, dan pelacakan status layanan.",
        "Mengintegrasikan modul Customer Satisfaction Index (CSI) untuk mengukur kepuasan pengguna setelah layanan selesai.",
        "Menyediakan dashboard petugas dan administrator untuk pengelolaan antrian, data, dan statistik pelayanan.",
    ])
    add_heading(doc, "1.4.2 Manfaat Penelitian", 3)
    add_body(doc, "Manfaat Teoritis: memperkaya literatur sistem informasi pelayanan publik dan penerapan CSI pada layanan kepolisian.")
    add_body(doc, "Manfaat Praktis:")
    add_bullets(doc, [
        "Bagi Polsek Palmerah: arsip digital terpusat dan transparansi status layanan.",
        "Bagi Petugas SPKT: antrian laporan digital (officer inbox), verifikasi, penugasan, dan pemrosesan terstruktur.",
        "Bagi Masyarakat: akses 24/7 untuk lapor, ajukan surat, pengaduan, dan lacak status tanpa login (track service).",
    ])

    add_heading(doc, "1.5 Batasan Masalah", 2)
    add_numbered(doc, [
        "Penelitian difokuskan pada perancangan sistem informasi SPKT berbasis web di Polsek Palmerah.",
        "Ruang lingkup meliputi modul: autentikasi, laporan polisi, layanan surat, pengaduan, informasi, notifikasi, dan CSI.",
        "Sistem menggunakan tiga role: masyarakat (user), petugas (petugas), dan administrator (admin).",
        "Database menggunakan SQLite (file data/spkt.db); file bukti disimpan di data/uploads/.",
        "Penilaian kepuasan menggunakan metode CSI dengan 5 dimensi penilaian (skala 1–4) terintegrasi dalam aplikasi.",
        "Di luar cakupan: integrasi email/SMS otomatis, verifikasi NIK nasional (Dukcapil), dan aplikasi mobile native.",
    ])

    add_heading(doc, "1.6 Sistematika Penulisan", 2)
    add_bullets(doc, [
        "BAB I Pendahuluan — latar belakang, masalah, tujuan, batasan.",
        "BAB II Tinjauan Pustaka — teori, literature review, analisis sistem berjalan, kebutuhan, desain sistem usulan.",
        "BAB III Metodologi — waktu/tempat, desain penelitian prototype, populasi/sampel, pengumpulan dan analisis data.",
        "BAB IV Hasil dan Pembahasan — implementasi SPKT Digital, pengujian, evaluasi CSI.",
        "BAB V Kesimpulan dan Saran.",
    ])
    doc.add_page_break()

    # === BAB II (key updated sections) ===
    add_heading(doc, "BAB II TINJAUAN PUSTAKA, KERANGKA KONSEPTUAL DAN HIPOTESIS", 1)
    add_body(doc, (
        "Bagian 2.1 Landasan Teori (definisi sistem informasi, pelayanan publik, SPKT, CSI, UML) "
        "dan 2.2 Literature Review mengacu pada naskah proposal asli serta daftar pustaka. "
        "Pada revisi ini, bagian analisis kebutuhan dan desain sistem (2.8–2.9) disesuaikan dengan "
        "implementasi aplikasi SPKT Digital yang telah dibangun."
    ))

    add_heading(doc, "2.2.2 Kebaruan Penelitian (Research Gap)", 2)
    add_body(doc, (
        "Penelitian sebelumnya sering mengukur CSI secara terpisah tanpa integrasi ke sistem informasi. "
        "Kebaruan penelitian ini: pengembangan SPKT Digital berbasis web (Next.js + SQLite) yang "
        "mengintegrasikan pelaporan online, layanan surat, pengaduan, pelacakan status, dashboard admin, "
        "dan modul CSI otomatis dengan 5 dimensi penilaian terkomputerisasi."
    ))

    add_heading(doc, "2.5 Hipotesis", 2)
    add_numbered(doc, [
        "Diduga sistem informasi SPKT berbasis web dapat mengubah pelaporan manual menjadi online.",
        "Diduga sistem memberikan kemudahan masyarakat melapor tanpa datang ke Polsek Palmerah.",
        "Diduga sistem meningkatkan efektivitas pengelolaan data laporan oleh petugas.",
        "Diduga modul CSI terintegrasi dapat mengukur kepuasan masyarakat secara objektif.",
    ])

    add_heading(doc, "2.6 Analisis Sistem Yang Sedang Berjalan", 2)
    add_body(doc, (
        "Berdasarkan observasi dan wawancara di SPKT Polsek Palmerah, sistem berjalan masih manual-digital campuran: "
        "masyarakat datang langsung, petugas input data ke komputer, pencetakan STPLK/LP manual, tanpa portal online "
        "untuk masyarakat. Tidak ada antrian elektronik, pelacakan status daring, maupun survei kepuasan terkomputerisasi."
    ))

    add_heading(doc, "2.7 Desain Sistem Yang Sedang Berjalan", 2)
    add_body(doc, "[Sisipkan Gambar 2.2 Activity Diagram Sistem Berjalan]")
    add_body(doc, "[Sisipkan Gambar 2.3 Use Case Diagram Sistem Berjalan]")

    add_heading(doc, "2.8 Analisis Kebutuhan Fungsional dan Non Fungsional", 2)
    add_heading(doc, "2.8.1 Kebutuhan Fungsional (Sesuai Aplikasi SPKT Digital)", 3)

    add_body(doc, "A. Modul Autentikasi dan Keamanan")
    add_bullets(doc, [
        "Registrasi dan login masyarakat, petugas, dan admin (email + password bcrypt).",
        "Sesi berbasis cookie HttpOnly (spkt_session) berlaku 7 hari.",
        "Logout, lupa password (reset dengan verifikasi email + NIK), autentikasi 2FA (TOTP/Google Authenticator).",
        "Pengaturan profil, ganti password, preferensi notifikasi dan privasi.",
        "Ekspor data pribadi (JSON) dan hapus akun dengan anonimisasi data terkait.",
    ])

    add_body(doc, "B. Modul Laporan Polisi (Masyarakat)")
    add_bullets(doc, [
        "Buat laporan kejadian (jenis, tanggal, lokasi, deskripsi) dengan upload bukti foto/dokumen.",
        "Simpan draft laporan dan lanjutkan pengisian dari menu Laporan Saya.",
        "Nomor referensi otomatis format LAP-YYYY-NNNN.",
        "Lihat daftar laporan milik sendiri (filter NIK) dengan timeline status.",
        "Status: draft → submitted → verified → assigned → processing → completed / rejected.",
    ])

    add_body(doc, "C. Modul Antrian dan Pemrosesan (Petugas)")
    add_bullets(doc, [
        "Dashboard antrian laporan masuk (officer inbox): belum ditugaskan + tugas petugas yang login.",
        "Verifikasi laporan, ambil tugas (assign), update proses, selesaikan atau tolak laporan.",
        "Proses layanan surat dan pengaduan dengan update status dan timeline.",
        "Notifikasi otomatis ke pelapor saat status berubah (sesuai preferensi user).",
    ])

    add_body(doc, "D. Modul Layanan Surat")
    add_bullets(doc, [
        "Pengajuan surat (SKCK, keterangan, dll.) dengan lampiran.",
        "Nomor SUR-YYYY-NNNN, timeline status, jadwal pengambilan.",
        "Unduh/cetak PDF surat saat status ready/completed.",
        "Status: draft → submitted → verified → ready → completed / rejected.",
    ])

    add_body(doc, "E. Modul Pengaduan")
    add_bullets(doc, [
        "Submit pengaduan dengan kategori, subjek, deskripsi, dan lampiran.",
        "Nomor PGD-YYYY-NNNN, timeline, tanggapan petugas.",
        "Status: submitted → reviewing → processing → resolved → closed.",
    ])

    add_body(doc, "F. Modul Lacak Layanan (Publik)")
    add_bullets(doc, [
        "Pelacakan laporan/surat/pengaduan tanpa login menggunakan nomor layanan + NIK.",
        "Menampilkan status dan timeline perkembangan.",
    ])

    add_body(doc, "G. Modul Administrator")
    add_bullets(doc, [
        "Kelola semua laporan dengan override status (tercatat audit log).",
        "Kelola user: suspend/aktifkan, ubah role.",
        "Kelola data petugas (CRUD officers, link ke akun login).",
        "Dashboard statistik, ringkasan CSI, kelola artikel informasi (CMS).",
        "Audit log aksi sensitif admin.",
    ])

    add_body(doc, "H. Modul CSI (Customer Satisfaction Index)")
    add_bullets(doc, [
        "Survei kepuasan setelah layanan selesai (laporan, surat, atau pengaduan).",
        "5 dimensi penilaian: Kemudahan Prosedur, Kecepatan Pelayanan, Keramahan Petugas, Kejelasan Informasi, Kualitas Hasil Layanan.",
        "Skala Likert 1–4 per dimensi; perhitungan skor CSI otomatis di backend.",
        "Dashboard admin menampilkan agregat kepuasan per jenis layanan.",
    ])

    add_body(doc, "I. Modul Pendukung")
    add_bullets(doc, [
        "Notifikasi in-app dengan polling.",
        "Halaman informasi/artikel edukasi.",
        "Pagination daftar data (laporan, surat, pengaduan).",
    ])

    add_heading(doc, "2.8.2 Kebutuhan Non Fungsional", 3)
    add_bullets(doc, [
        "Keamanan: bcrypt password, cookie HttpOnly, rate limit login/track/forgot-password, RBAC 3 role.",
        "Usability: antarmuka responsif (desktop + mobile), Tailwind CSS + Radix UI.",
        "Performance: pagination 20 item/halaman, SQLite WAL mode.",
        "Portability: web-based, akses via browser modern.",
        "Maintainability: arsitektur 3 lapis (React FE → Next.js API Routes → Service Layer → SQLite).",
        "Reliability: healthcheck /api/health, backup database script (npm run backup:db).",
        "Runtime: Node.js ≥ 22.5 (node:sqlite native).",
    ])

    add_heading(doc, "2.9 Desain Sistem Usulan", 2)
    add_heading(doc, "2.9.1 Arsitektur Teknis SPKT Digital", 3)
    add_body(doc, (
        "Sistem dikembangkan sebagai Single Page Application (SPA) dalam satu proyek Next.js 15:\n"
        "• Frontend: React 19, komponen di src/components/, routing internal ?view= per role.\n"
        "• Client API: src/lib/spktApi.ts (fetch ke /api/* dengan credentials cookie).\n"
        "• Backend: Next.js Route Handlers di src/app/api/ (~49 endpoint REST).\n"
        "• Business Logic: src/lib/services/ (spkt.ts, users.ts, notifications.ts, audit.ts, csi.ts).\n"
        "• Database: SQLite file data/spkt.db (21 tabel), upload fisik data/uploads/."
    ))

    add_heading(doc, "2.9.2 Entity Relationship Diagram (ERD)", 3)
    add_body(doc, "[Sisipkan Gambar 2.4 ERD — entitas utama:]")
    add_bullets(doc, [
        "users (pusat) → sessions, notifications, user_activities, officers, pending_2fa",
        "reports → report_timeline, report_evidence; assigned_officer_id → officers",
        "letter_requests → letter_timeline, letter_attachments",
        "complaints → complaint_timeline, complaint_files",
        "satisfaction_surveys → survey_responses → survey_dimensions",
        "Pendukung: reference_counters, audit_logs, info_articles",
    ])

    add_heading(doc, "2.9.3 Use Case Diagram Sistem Usulan", 3)
    add_body(doc, "[Sisipkan Gambar 2.6 Use Case Diagram]")
    add_body(doc, "Aktor dan use case utama:")
    add_bullets(doc, [
        "Masyarakat: register, login, buat laporan, draft laporan, ajukan surat, pengaduan, lacak layanan, isi CSI, pengaturan akun.",
        "Petugas: login, lihat antrian laporan, verifikasi, assign, proses laporan/surat/pengaduan, update status.",
        "Admin: semua fitur petugas + kelola user, kelola petugas, statistik, CSI dashboard, kelola artikel, audit log.",
        "Pengunjung anonim: lacak layanan, baca artikel informasi.",
    ])

    add_heading(doc, "2.9.4 Activity Diagram Sistem Usulan", 3)
    add_body(doc, "[Sisipkan Gambar 2.5 Activity Diagram — alur: login → buat laporan → verifikasi petugas → selesai → CSI]")

    add_heading(doc, "2.9.5 Class Diagram dan CRC Card", 3)
    add_body(doc, "[Sisipkan Gambar 2.7 Class Diagram dan Tabel 2.8 CRC Card]")
    add_body(doc, "Kelas utama: User (Masyarakat, Petugas, Admin), Report, LetterRequest, Complaint, Officer, Notification, SatisfactionSurvey, AuditLog.")

    add_heading(doc, "2.9.6 Alur Status Layanan", 3)
    add_body(doc, "Laporan: draft → submitted → verified → assigned → processing → completed | rejected")
    add_body(doc, "Surat: draft → submitted → verified → ready → completed | rejected")
    add_body(doc, "Pengaduan: submitted → reviewing → processing → resolved → closed")

    doc.add_page_break()

    # === BAB III updates ===
    add_heading(doc, "BAB III METODOLOGI PENELITIAN", 1)
    add_heading(doc, "3.1 Waktu dan Tempat Penelitian", 2)
    add_body(doc, "Tempat: SPKT Polsek Palmerah, Jalan Palmerah Barat III No. 1, Palmerah, Jakarta Barat.")
    add_body(doc, "(Bagian sejarah instansi dan struktur organisasi mengacu proposal asli — Gambar 3.1)")

    add_heading(doc, "3.2 Desain Penelitian", 2)
    add_body(doc, "Metode Prototype: identifikasi masalah → pengumpulan data → perancangan → pembangunan SPKT Digital → pengujian → evaluasi CSI → kesimpulan.")
    add_body(doc, "[Sisipkan Gambar 3.2 Flowchart Desain Penelitian]")

    add_heading(doc, "3.4 Teknik Pengumpulan Data — Instrumen CSI (Disesuaikan Aplikasi)", 2)
    add_body(doc, (
        "Survei CSI diimplementasikan dalam aplikasi (bukan kuesioner kertas terpisah). "
        "Setelah layanan selesai, pengguna mengisi 5 dimensi dengan skala 1–4:"
    ))
    add_numbered(doc, [
        "Kemudahan Prosedur (ease)",
        "Kecepatan Pelayanan (speed)",
        "Keramahan Petugas (officer)",
        "Kejelasan Informasi (clarity)",
        "Kualitas Hasil Layanan (quality)",
    ])
    add_body(doc, "Skala: 1 = Sangat Tidak Puas, 2 = Tidak Puas, 3 = Puas, 4 = Sangat Puas.")
    add_body(doc, "Rumus CSI di sistem: rata-rata tertimbang skor dimensi / skor maksimum (4) × 100%.")

    add_heading(doc, "3.5 Stack Pengembangan", 2)
    add_bullets(doc, [
        "Framework: Next.js 15 (App Router)",
        "Frontend: React 19, TypeScript, Tailwind CSS, Radix UI",
        "Backend: Next.js API Routes, Node.js ≥ 22.5",
        "Database: SQLite (node:sqlite), file data/spkt.db",
        "Testing: Vitest",
        "Editor: Visual Studio Code / Cursor IDE",
    ])

    add_heading(doc, "3.6 Teknik Analisis Data", 2)
    add_bullets(doc, [
        "Analisis deskriptif kualitatif (observasi, wawancara).",
        "Analisis fungsional sistem (black-box testing fitur per role).",
        "Analisis CSI dari data satisfaction_surveys dan survey_responses di database.",
    ])

    add_body(doc, "(Bagian jadwal penelitian Tabel 3.2 dan daftar pustaka mengacu proposal asli.)")

    doc.add_page_break()

    # === Lampiran teknis ===
    add_heading(doc, "LAMPIRAN — SPESIFIKASI TEKNIS APLIKASI", 1)
    add_heading(doc, "Akun Demo Sistem", 2)
    add_body(doc, "Masyarakat: user@spkt.id / spkt123")
    add_body(doc, "Petugas: petugas@spkt.id / spkt123")
    add_body(doc, "Admin: admin@spkt.id / spkt123")

    add_heading(doc, "Endpoint API Utama", 2)
    add_bullets(doc, [
        "POST /api/auth/login, GET /api/auth/session",
        "GET/POST /api/reports, PATCH /api/reports/:id",
        "GET/POST /api/letters, GET /api/letters/:id/pdf",
        "GET/POST /api/complaints",
        "GET /api/track (publik)",
        "POST /api/survey/submit, GET /api/survey/csi/summary",
        "GET /api/stats/admin, GET /api/audit-logs",
    ])

    add_body(doc, "Dokumentasi lengkap: docs/presentasi-ppt-teknis.md dan README.md pada repositori proyek.")

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    build()
